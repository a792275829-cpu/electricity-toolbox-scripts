from __future__ import annotations

import re
import unittest

from docx import Document

from generate_red_marked_report import (
    LOCAL_TEMPLATE_PATH,
    QINGYUAN_DAILY_CLEARING_VALUES,
    ensure_qingyuan_template_content,
    startup_date_paragraphs,
)


class QingyuanTemplateTests(unittest.TestCase):
    def test_qingyuan_content_is_complete_and_idempotent(self) -> None:
        document = Document(LOCAL_TEMPLATE_PATH)
        ensure_qingyuan_template_content(document)
        ensure_qingyuan_template_content(document)

        dates = startup_date_paragraphs(document)
        self.assertGreaterEqual(len(dates), 2)
        for date_paragraph in dates[:2]:
            start = next(
                index
                for index, paragraph in enumerate(document.paragraphs)
                if paragraph._element is date_paragraph._element
            ) + 1
            following = document.paragraphs[start:]
            next_boundary = next(
                (
                    index
                    for index, paragraph in enumerate(following)
                    if re.fullmatch(r"(?:\d+月)?\d+日：", paragraph.text.strip())
                    or paragraph.text.strip() == "二、分公司电量、电价完成情况"
                ),
                len(following),
            )
            block = following[:next_boundary]
            self.assertEqual(
                sum(paragraph.text.strip().startswith("清远厂：") for paragraph in block),
                1,
            )

        qingyuan_rows = [
            row
            for row in document.tables[0].rows
            if re.sub(r"\s+", "", row.cells[0].text) == "清远"
        ]
        self.assertEqual(len(qingyuan_rows), 1)
        self.assertEqual(
            [cell.text.strip() for cell in qingyuan_rows[0].cells],
            QINGYUAN_DAILY_CLEARING_VALUES,
        )


if __name__ == "__main__":
    unittest.main()
