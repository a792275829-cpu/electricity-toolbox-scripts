#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import sys
import time
import urllib.parse
import urllib.request
from copy import deepcopy
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

from openpyxl import load_workbook
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

from fetch_daily_data import build_report_date_plan, run_daily_data


SCRIPT_DIR = Path(__file__).resolve().parent
APP_ROOT = SCRIPT_DIR.parent if SCRIPT_DIR.name in {"scripts", "\u811a\u672c"} else SCRIPT_DIR
OUTPUT_DIR = APP_ROOT / "\u8f93\u51fa"
CLEARING_WORKBOOK_DIR = Path.home() / "Downloads"
LOCAL_TEMPLATE_PATH = APP_ROOT / "\u6a21\u677f" / "\u6bcf\u65e5\u751f\u4ea7\u7ecf\u8425\u60c5\u51b5\u6c47\u62a5\uff08\u81ea\u52a8\uff09.docx"
TEMPLATE_PATH = LOCAL_TEMPLATE_PATH if LOCAL_TEMPLATE_PATH.exists() else (
    Path.home()
    / "Desktop"
    / "\u590d\u76d8"
    / "\u6bcf\u65e5\u751f\u4ea7\u7ecf\u8425\u60c5\u51b5\u6c47\u62a5\uff08\u81ea\u52a8\uff09.docx"
)


LOAD_SUBJECT_CODES = ["1", "6", "2", "18"]
GUANGDONG_WEATHER_LATITUDE = 23.1291
GUANGDONG_WEATHER_LONGITUDE = 113.2644


def log(message: str) -> None:
    print(message, flush=True)


def date_label(date_text: str) -> str:
    dt = datetime.strptime(date_text, "%Y-%m-%d")
    return f"{dt.month}\u6708{dt.day}\u65e5"


def day_number(date_text: str) -> str:
    return str(datetime.strptime(date_text, "%Y-%m-%d").day)


def report_filename(date_text: str) -> str:
    dt = datetime.strptime(date_text, "%Y-%m-%d")
    return f"\u6bcf\u65e5\u751f\u4ea7\u7ecf\u8425\u60c5\u51b5\u6c47\u62a5\uff08{dt.month}\u6708{dt.day}\u65e5\uff09.docx"


def previous_report_path(report_date: str, output_dir: Path) -> Path:
    previous_day = datetime.strptime(report_date, "%Y-%m-%d").date() - timedelta(days=1)
    return output_dir / report_filename(previous_day.strftime("%Y-%m-%d"))


def resolve_template_path(
    report_date: str,
    template_path: Path,
    output_dir: Path,
    *,
    prefer_previous_report: bool = True,
) -> Path:
    previous_path = previous_report_path(report_date, output_dir)
    if prefer_previous_report and previous_path.exists():
        return previous_path
    return template_path


def validate_date(date_text: str) -> str:
    try:
        return datetime.strptime(date_text, "%Y-%m-%d").strftime("%Y-%m-%d")
    except ValueError as exc:
        raise ValueError(f"invalid date: {date_text}, expected YYYY-MM-DD") from exc


def trend_word(value: float) -> str:
    return "\u4e0a\u5347" if value >= 0 else "\u4e0b\u964d"


def change_word(value: float) -> str:
    return "\u589e\u52a0" if value >= 0 else "\u51cf\u5c11"


def mw_to_10k_kw(value_mw: float) -> int:
    return round(value_mw / 10)


def fmt_number(value: float | None, digits: int = 0) -> str:
    if value is None:
        return "\u5f85\u8865\u5145"
    if digits == 0:
        return str(int(round(value)))
    return f"{value:.{digits}f}".rstrip("0").rstrip(".")


def weather_code_label(code: Any) -> str:
    try:
        value = int(code)
    except (TypeError, ValueError):
        return "\u591a\u4e91"
    if value == 0:
        return "\u6674"
    if value in {1, 2}:
        return "\u591a\u4e91"
    if value == 3:
        return "\u9634"
    if value in {45, 48}:
        return "\u96fe"
    if value in {51, 53, 55, 56, 57, 61, 80}:
        return "\u5c0f\u96e8"
    if value in {63, 66, 81}:
        return "\u4e2d\u96e8"
    if value in {65, 67, 82}:
        return "\u5927\u96e8"
    if value in {71, 73, 75, 77, 85, 86}:
        return "\u96ea"
    if value in {95, 96, 99}:
        return "\u96f7\u9635\u96e8"
    return "\u591a\u4e91"


def fetch_guangdong_weather_text(report_date: str) -> str:
    start = datetime.strptime(report_date, "%Y-%m-%d").date()
    end = start + timedelta(days=2)
    query = urllib.parse.urlencode(
        {
            "latitude": GUANGDONG_WEATHER_LATITUDE,
            "longitude": GUANGDONG_WEATHER_LONGITUDE,
            "daily": "weather_code,temperature_2m_max,temperature_2m_min",
            "timezone": "Asia/Shanghai",
            "start_date": start.strftime("%Y-%m-%d"),
            "end_date": end.strftime("%Y-%m-%d"),
        }
    )
    url = f"https://api.open-meteo.com/v1/forecast?{query}"
    last_error: Exception | None = None
    for _ in range(3):
        try:
            with urllib.request.urlopen(url, timeout=30) as response:
                data = json.load(response)
            break
        except Exception as exc:
            last_error = exc
            time.sleep(2)
    else:
        raise RuntimeError(f"\u8bfb\u53d6\u5e7f\u4e1c\u5929\u6c14\u5931\u8d25\uff1a{last_error}") from last_error
    daily = data.get("daily") or {}
    labels = [weather_code_label(code) for code in daily.get("weather_code") or []]
    max_values = [parse_float(value) for value in daily.get("temperature_2m_max") or [] if value is not None]
    min_values = [parse_float(value) for value in daily.get("temperature_2m_min") or [] if value is not None]
    if not labels or not max_values or not min_values:
        raise RuntimeError("\u672a\u8bfb\u53d6\u5230\u5e7f\u4e1c\u672a\u6765\u4e09\u5929\u5929\u6c14\u6570\u636e")
    weather_text = labels[0] if labels[0] == labels[-1] else f"{labels[0]}\u8f6c{labels[-1]}"
    return f"\u672a\u6765\u4e09\u5929{weather_text}\uff0c\u6c14\u6e29{round(min(min_values))}\u2103-{round(max(max_values))}\u2103\uff0c"


def replace_weather_prefix(paragraph, report_date: str) -> None:
    prefix = fetch_guangdong_weather_text(report_date)
    text = paragraph.text
    updated = re.sub(r"^\u672a\u6765\u4e09\u5929.*?\uff0c\u6c14\u6e29\s*.*?\u2103\s*-\s*.*?\u2103\uff0c", prefix, text, count=1)
    if updated == text:
        updated = prefix + text
    replace_paragraph(paragraph, updated)


def parse_float(value: Any) -> float:
    if value is None:
        return 0.0
    return float(value)


def find_clearing_workbook(date_text: str, workbook_dir: Path = CLEARING_WORKBOOK_DIR) -> Path:
    dt = datetime.strptime(date_text, "%Y-%m-%d")
    date_token = f"{dt.year}.{dt.month}.{dt.day}"
    candidates = [
        path
        for path in workbook_dir.glob("*.xlsx")
        if not path.name.startswith("~$")
        and date_token in path.name
        and "\u51fa\u6e05\u60c5\u51b5" in path.name
    ]
    if not candidates:
        raise FileNotFoundError(f"未找到 {date_token} 的出清情况 Excel：{workbook_dir}")
    return max(candidates, key=lambda path: path.stat().st_mtime)


def load_day_ahead_clearing_summary(date_text: str, workbook_path: Path | None = None) -> dict[str, Any]:
    if workbook_path is None:
        workbook_path = find_clearing_workbook(date_text)
    workbook = load_workbook(workbook_path, read_only=True, data_only=True)
    sheet_name = next(
        (
            name
            for name in workbook.sheetnames
            if "\u5feb\u62a5" in name and "\u65e5\u524d" in name
        ),
        None,
    )
    if sheet_name is None:
        raise KeyError(f"{workbook_path} 中未找到快报（日前）工作表")
    sheet = workbook[sheet_name]

    rows = {
        "company": 7,
        "\u6c55\u5934": 8,
        "\u6d77\u95e8": 13,
        "\u4e1c\u839e": 19,
        "\u6d77\u4e0a\u98ce\u7535": 29,
        "\u9b80\u83b2": 30,
        "\u5f52\u6e56": 31,
    }
    return {
        "workbook": str(workbook_path),
        "sheet": sheet_name,
        "totalEnergy": parse_float(sheet.cell(rows["company"], 6).value),
        "averagePrice": parse_float(sheet.cell(rows["company"], 7).value),
        "companies": {
            "\u6c55\u5934": {
                "energy": parse_float(sheet.cell(rows["\u6c55\u5934"], 6).value),
                "price": parse_float(sheet.cell(rows["\u6c55\u5934"], 7).value),
            },
            "\u6d77\u95e8": {
                "energy": parse_float(sheet.cell(rows["\u6d77\u95e8"], 6).value),
                "price": parse_float(sheet.cell(rows["\u6d77\u95e8"], 7).value),
            },
            "\u4e1c\u839e": {
                "energy": parse_float(sheet.cell(rows["\u4e1c\u839e"], 6).value),
                "price": parse_float(sheet.cell(rows["\u4e1c\u839e"], 7).value),
            },
            "\u6d77\u4e0a\u98ce\u7535": {
                "energy": parse_float(sheet.cell(rows["\u6d77\u4e0a\u98ce\u7535"], 7).value),
                "price": parse_float(sheet.cell(rows["\u6d77\u4e0a\u98ce\u7535"], 8).value),
            },
            "\u9b80\u83b2": {
                "energy": parse_float(sheet.cell(rows["\u9b80\u83b2"], 7).value),
                "price": parse_float(sheet.cell(rows["\u9b80\u83b2"], 8).value),
            },
            "\u5f52\u6e56": {
                "energy": parse_float(sheet.cell(rows["\u5f52\u6e56"], 7).value),
                "price": parse_float(sheet.cell(rows["\u5f52\u6e56"], 8).value),
            },
        },
    }


def load_online_price_summary(date_text: str, workbook_path: Path | None = None) -> dict[str, Any]:
    if workbook_path is None:
        workbook_path = find_clearing_workbook(date_text)
    workbook = load_workbook(workbook_path, read_only=True, data_only=True)
    sheet_name = next(
        (
            name
            for name in workbook.sheetnames
            if "\u5feb\u62a5" in name and "\u65e5\u524d" in name
        ),
        None,
    )
    if sheet_name is None:
        raise KeyError(f"{workbook_path} 中未找到快报（日前）工作表")
    sheet = workbook[sheet_name]
    return {
        "workbook": str(workbook_path),
        "sheet": sheet_name,
        "averagePrice": parse_float(sheet["G7"].value),
        "companies": {
            "\u6c55\u5934": parse_float(sheet["G8"].value),
            "\u6d77\u95e8": parse_float(sheet["G13"].value),
            "\u4e1c\u839e": parse_float(sheet["G19"].value),
            "\u6d77\u4e0a\u98ce\u7535": parse_float(sheet["H29"].value),
            "\u9b80\u83b2": parse_float(sheet["H30"].value),
            "\u5f52\u6e56": parse_float(sheet["H31"].value),
        },
    }


def copy_run_format(source, target) -> None:
    source_rpr = source._element.rPr
    if source_rpr is not None:
        target_rpr = target._element.rPr
        if target_rpr is not None:
            target._element.remove(target_rpr)
        target._element.insert(0, deepcopy(source_rpr))
    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    target.font.name = source.font.name
    target.font.size = source.font.size

    # Word uses w:eastAsia for Chinese text. python-docx's font.name alone only
    # sets ascii/hAnsi in some cases, which causes visible font fallback.
    rpr = target._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    font_name = source.font.name or "\u4eff\u5b8b"
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def replace_paragraph(paragraph, text: str, *, keep_red: bool = False) -> None:
    template_run = paragraph.runs[0] if paragraph.runs else None
    for run_idx in range(len(paragraph.runs) - 1, -1, -1):
        paragraph._element.remove(paragraph.runs[run_idx]._element)
    run = paragraph.add_run(text)
    if template_run is not None:
        copy_run_format(template_run, run)
    if not keep_red:
        run.font.color.rgb = RGBColor(0, 0, 0)


def replace_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph(paragraph, text)


def replace_red_runs(paragraph, values: list[str]) -> None:
    value_iter = iter(values)
    for run in paragraph.runs:
        color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
        if color and str(color).upper() == "FF0000":
            try:
                run.text = next(value_iter)
            except StopIteration:
                return


def clear_paragraph_red(paragraph) -> None:
    for run in paragraph.runs:
        color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
        if color and str(color).upper() == "FF0000":
            run.font.color.rgb = RGBColor(0, 0, 0)


def replace_leading_date(paragraph, date_text: str) -> None:
    text = paragraph.text
    updated = re.sub(r"^\d+月\d+日", date_label(date_text), text, count=1)
    if updated != text:
        replace_paragraph(paragraph, updated)


def get_company_energy(run_date: str) -> dict[str, Any]:
    rows = collect_detail_rows(run_date, load_config())
    totals: dict[str, float] = {}
    for row in rows:
        company = str(row.get("company") or "")
        label = None
        for raw, display in {
            "\u6c55\u5934\u7535\u5382": "\u6c55\u5934",
            "\u6d77\u95e8\u7535\u5382": "\u6d77\u95e8",
            "\u8c22\u5c97\u7535\u5382": "\u4e1c\u839e",
            "\u534e\u80fd\u5e7f\u4e1c\u6c55\u5934\u6d77\u4e0a\u98ce\u7535\u6709\u9650\u8d23\u4efb\u516c\u53f8": "\u6d77\u98ce",
            "\u534e\u80fd\uff08\u6c55\u5934\u91d1\u5e73\uff09\u65b0\u80fd\u6e90\u6709\u9650\u8d23\u4efb\u516c\u53f8": "\u9b80\u83b2",
            "\u534e\u80fd\uff08\u6f6e\u5dde\u6f6e\u5b89\uff09\u65b0\u80fd\u6e90\u6709\u9650\u8d23\u4efb\u516c\u53f8": "\u5f52\u6e56",
        }.items():
            if raw == company:
                label = display
                break
        if label is None:
            label = company
        totals[label] = totals.get(label, 0.0) + float(row.get("onlineEle") or 0.0)
    return {
        "total": round(sum(totals.values())),
        "companies": {name: round(totals.get(name, 0.0)) for name in ["\u6c55\u5934", "\u6d77\u95e8", "\u4e1c\u839e", "\u6d77\u98ce", "\u9b80\u83b2", "\u5f52\u6e56"]},
    }


def metric(subject: dict[str, Any], name: str) -> dict[str, Any]:
    return subject["comparison"][name]


def subject_by_code(payload: dict[str, Any], code: str) -> dict[str, Any]:
    for subject in payload["loadCompare"]["subjects"]:
        if str(subject.get("loadType")) == code:
            return subject
    raise KeyError(code)


def build_strategy_text(payload: dict[str, Any]) -> str:
    load = subject_by_code(payload, "1")
    a_power = subject_by_code(payload, "2")
    local = subject_by_code(payload, "18")
    price = next(row for row in payload["priceBoard"]["rows"] if row["dataType"] == "PRE")
    public_data = payload.get("operationPublicData") or {}
    capacity = (public_data.get("capacity") or {}).get("selected", {}).get("capacityWanKw")
    startup = (public_data.get("startupCount") or {}).get("selected", {}).get("maxStartupCount")
    startup_diff = (public_data.get("startupCount") or {}).get("difference")
    load_max = metric(load, "\u6700\u5927\u503c(MW)")
    load_min = metric(load, "\u6700\u5c0f\u503c(MW)")
    a_avg = metric(a_power, "\u5e73\u5747\u503c(MW)")
    local_avg = metric(local, "\u5e73\u5747\u503c(MW)")
    return (
        f"\u8fd0\u884c\u65e5\u7edf\u8c03\u8d1f\u8377\u9700\u6c42{trend_word(load_max['difference'])}\uff0c"
        f"\u5168\u5929\u6700\u5927{mw_to_10k_kw(load_max['selected'])}\u4e07\u5343\u74e6\uff0c"
        f"\u73af\u6bd4{trend_word(load_max['difference'])}{abs(mw_to_10k_kw(load_max['difference']))}\u4e07\u5343\u74e6\uff0c"
        f"\u6700\u5c0f\u503c{mw_to_10k_kw(load_min['selected'])}\u4e07\u5343\u74e6\uff0c"
        f"\u73af\u6bd4{trend_word(load_min['difference'])}{abs(mw_to_10k_kw(load_min['difference']))}\u4e07\u5343\u74e6\uff1b"
        f"A\u7c7b\u7535\u6e90\u51fa\u529b{change_word(a_avg['difference'])}\u3001"
        f"\u5730\u65b9\u7535\u6e90\u51fa\u529b{change_word(local_avg['difference'])}\u3002"
        f"\u68c0\u4fee\u5bb9\u91cf{fmt_number(capacity)}\u4e07\u5343\u74e6\u3002"
        f"\u5e02\u573a\u5f00\u673a\u53f0\u6570{fmt_number(startup)}\u53f0\uff0c"
        f"\u73af\u6bd4{change_word(startup_diff or 0)}{fmt_number(abs(startup_diff or 0))}\u53f0\uff0c"
        f"\u5168\u7701\u65e5\u524d\u51fa\u6e05\u5747\u4ef7{fmt_number(price['selectedDailyAvgPrice'], 2)}\u5143/\u5146\u74e6\u65f6\uff0c"
        f"\u73af\u6bd4{trend_word(price['difference'] or 0)}{fmt_number(abs(price['difference'] or 0), 2)}\u5143/\u5146\u74e6\u65f6\u3002"
    )


def build_realtime_text(payload: dict[str, Any]) -> str:
    load = subject_by_code(payload, "1")
    b_space = subject_by_code(payload, "6")
    a_power = subject_by_code(payload, "2")
    local = subject_by_code(payload, "18")
    price = next(row for row in payload["priceBoard"]["rows"] if row["dataType"] == "RT")
    load_max = metric(load, "\u6700\u5927\u503c(MW)")
    load_min = metric(load, "\u6700\u5c0f\u503c(MW)")
    b_max = metric(b_space, "\u6700\u5927\u503c(MW)")
    b_min = metric(b_space, "\u6700\u5c0f\u503c(MW)")
    a_avg = metric(a_power, "\u5e73\u5747\u503c(MW)")
    local_avg = metric(local, "\u5e73\u5747\u503c(MW)")
    return (
        f"\u5b9e\u65f6\u7edf\u8c03\u8d1f\u8377\u9700\u6c42{trend_word(load_max['difference'])}\uff0c"
        f"\u5168\u5929\u6700\u5927\u503c{mw_to_10k_kw(load_max['selected'])}\u4e07\u5343\u74e6\uff0c"
        f"\u6700\u5c0f\u503c{mw_to_10k_kw(load_min['selected'])}\u4e07\u5343\u74e6\uff0c"
        f"A\u7c7b\u7535\u51fa\u529b{change_word(a_avg['difference'])}\u3001"
        f"\u5730\u65b9\u7535\u51fa\u529b{change_word(local_avg['difference'])}\uff0c"
        f"B\u7c7b\u7ade\u4ef7\u7a7a\u95f4\u5168\u5929\u6700\u5927\u503c{mw_to_10k_kw(b_max['selected'])}\u4e07\u5343\u74e6\uff0c"
        f"\u6700\u5c0f\u503c{mw_to_10k_kw(b_min['selected'])}\u4e07\u5343\u74e6\uff0c"
        f"\u5168\u5929\u5b9e\u65f6\u5747\u4ef7{fmt_number(price['selectedDailyAvgPrice'], 2)}\u5143/\u5146\u74e6\u65f6\u3002"
    )


def load_info_item(market_data: dict[str, Any], load_type: str, kind: str) -> dict[str, Any]:
    return market_data["loadInfo"]["byType"][load_type][kind]


def market_price_item(market_data: dict[str, Any], data_type: str) -> dict[str, Any]:
    return next(row for row in market_data["priceBoard"]["rows"] if row["dataType"] == data_type)


def build_day_ahead_market_text(payload: dict[str, Any]) -> str:
    market_data = payload["marketAnalysisData"]
    load = load_info_item(market_data, "1", "dayAhead")
    a_power = load_info_item(market_data, "2", "dayAhead")
    local = load_info_item(market_data, "18", "dayAhead")
    price = market_price_item(market_data, "PRE")
    public_data = market_data.get("operationPublicData") or {}
    capacity = (public_data.get("capacity") or {}).get("selected", {}).get("capacityWanKw")
    startup = (public_data.get("startupCount") or {}).get("selected", {}).get("maxStartupCount")
    startup_diff = (public_data.get("startupCount") or {}).get("difference")
    selected = load["selected"]
    return (
        f"\u8fd0\u884c\u65e5\u7edf\u8c03\u8d1f\u8377\u9700\u6c42{trend_word(load.get('maxDiff') or 0)}\uff0c"
        f"\u5168\u5929\u6700\u5927\u503c{mw_to_10k_kw(selected.get('max'))}\u4e07\u5343\u74e6\uff0c"
        f"\u73af\u6bd4{trend_word(load.get('maxDiff') or 0)}{abs(mw_to_10k_kw(load.get('maxDiff') or 0))}\u4e07\u5343\u74e6\uff0c"
        f"\u6700\u5c0f\u503c{mw_to_10k_kw(selected.get('min'))}\u4e07\uff0c"
        f"\u73af\u6bd4{trend_word(load.get('minDiff') or 0)}{abs(mw_to_10k_kw(load.get('minDiff') or 0))}\u4e07\u5343\u74e6\uff1b"
        f"A\u7c7b\u7535\u51fa\u529b{change_word(a_power.get('avgDiff') or 0)}\u3001"
        f"\u5730\u65b9\u7535\u51fa\u529b{change_word(local.get('avgDiff') or 0)}\uff0c"
        f"\u68c0\u4fee\u5bb9\u91cf{fmt_number(capacity)}\u4e07\u5343\u74e6\u3002"
        f"\u5e02\u573a\u5f00\u673a\u53f0\u6570{fmt_number(startup)}\u53f0\uff0c"
        f"\u73af\u6bd4{change_word(startup_diff or 0)}{fmt_number(abs(startup_diff or 0))}\u53f0\uff0c"
        f"\u5168\u7701\u65e5\u524d\u51fa\u6e05\u5747\u4ef7{fmt_number(price['selectedDailyAvgPrice'], 2)}\u5143/\u5146\u74e6\u65f6\uff0c"
        f"\u73af\u6bd4{trend_word(price['difference'] or 0)}{fmt_number(abs(price['difference'] or 0), 2)}\u5143/\u5146\u74e6\u65f6\u3002"
    )


def build_actual_market_text(payload: dict[str, Any]) -> str:
    market_data = payload["marketAnalysisData"]
    load = load_info_item(market_data, "1", "actual")
    a_power = load_info_item(market_data, "2", "actual")
    local = load_info_item(market_data, "18", "actual")
    b_space = load_info_item(market_data, "6", "actual")
    price = market_price_item(market_data, "RT")
    selected = load["selected"]
    b_selected = b_space["selected"]
    return (
        f"\u5b9e\u65f6\u7edf\u8c03\u8d1f\u8377\u9700\u6c42{trend_word(load.get('maxDiff') or 0)}\uff0c"
        f"\u5168\u5929\u6700\u5927\u503c{mw_to_10k_kw(selected.get('max'))}\u4e07\u5343\u74e6\uff0c"
        f"\u6700\u5c0f\u503c{mw_to_10k_kw(selected.get('min'))}\u4e07\u5343\u74e6\uff0c"
        f"A\u7c7b\u7535\u51fa\u529b{change_word(a_power.get('avgDiff') or 0)}\u3001"
        f"\u5730\u65b9\u7535\u51fa\u529b{change_word(local.get('avgDiff') or 0)}\uff0c"
        f"B\u7c7b\u7ade\u4ef7\u7a7a\u95f4\u5168\u5929\u6700\u5927\u503c{mw_to_10k_kw(b_selected.get('max'))}\u4e07\u5343\u74e6\uff0c"
        f"\u6700\u5c0f\u503c{mw_to_10k_kw(b_selected.get('min'))}\u4e07\u5343\u74e6\uff0c"
        f"\u5168\u5929\u5b9e\u65f6\u5747\u4ef7{fmt_number(price['selectedDailyAvgPrice'], 2)}\u5143/\u5146\u74e6\u65f6\u3002"
    )


def build_energy_text(date_text: str, energy: dict[str, Any], kind: str) -> str:
    labels = [
        ("\u6c55\u5934", "\u6c55\u5934"),
        ("\u6d77\u95e8", "\u6d77\u95e8"),
        ("\u4e1c\u839e", "\u4e1c\u839e"),
        ("\u6d77\u98ce", "\u6d77\u4e0a\u98ce\u7535"),
        ("\u9b80\u83b2", "\u9b80\u83b2"),
        ("\u5f52\u6e56", "\u5f52\u6e56"),
    ]
    parts = [f"{label}{int(energy['companies'].get(key, 0))}\u4e07" for key, label in labels]
    if kind == "online":
        return (
            f"{date_label(date_text)}\u5206\u516c\u53f8\u4e0a\u7f51\u7535\u91cf{int(energy['total'])}\u4e07\u5343\u74e6\u65f6\uff0c"
            f"\u5176\u4e2d\uff0c{'、'.join(parts)}\u5343\u74e6\u65f6\uff1b"
            f"\u5206\u516c\u53f8\u65e5\u524d\u5747\u4ef7\u5f85\u8865\u5145\u5398/\u5343\u74e6\u65f6\uff0c"
            f"\u5176\u4e2d\u6c55\u5934\u5f85\u8865\u5145\u5398\u3001\u6d77\u95e8\u5f85\u8865\u5145\u5398\u3001\u4e1c\u839e\u5f85\u8865\u5145\u5398\u3001"
            f"\u6d77\u4e0a\u98ce\u7535\u5f85\u8865\u5145\u5398\u3001\u9b80\u83b2\u5f85\u8865\u5145\u5398\u3001\u5f52\u6e56\u5f85\u8865\u5145\u5398/\u5343\u74e6\u65f6\u3002"
        )
    return (
        f"{date_label(date_text)}\u5206\u516c\u53f8\u65e5\u524d\u51fa\u6e05\u7535\u91cf{int(energy['total'])}\u4e07\u5343\u74e6\u65f6\u3002"
        f"\u5176\u4e2d\uff0c{'、'.join(parts)}\u5343\u74e6\u65f6\uff1b"
        f"\u5206\u516c\u53f8\u65e5\u524d\u5747\u4ef7\u5f85\u8865\u5145\u5398/\u5343\u74e6\u65f6\uff0c"
        f"\u5176\u4e2d\u6c55\u5934\u5f85\u8865\u5145\u5398\u3001\u6d77\u95e8\u5f85\u8865\u5145\u5398\u3001\u4e1c\u839e\u5f85\u8865\u5145\u5398\u3001"
        f"\u6d77\u4e0a\u98ce\u7535\u5f85\u8865\u5145\u5398\u3001\u9b80\u83b2\u5f85\u8865\u5145\u5398\u3001\u5f52\u6e56\u5f85\u8865\u5145\u5398/\u5343\u74e6\u65f6\u3002"
    )


def build_day_ahead_clearing_text(date_text: str, summary: dict[str, Any]) -> str:
    companies = summary["companies"]
    shantou = companies["\u6c55\u5934"]
    haimen = companies["\u6d77\u95e8"]
    dongguan = companies["\u4e1c\u839e"]
    offshore = companies["\u6d77\u4e0a\u98ce\u7535"]
    tuolian = companies["\u9b80\u83b2"]
    guihu = companies["\u5f52\u6e56"]
    return (
        f"{date_label(date_text)}\u5206\u516c\u53f8\u65e5\u524d\u51fa\u6e05\u7535\u91cf{fmt_number(summary['totalEnergy'])}\u4e07\u5343\u74e6\u65f6\u3002"
        f"\u5176\u4e2d\uff0c\u6c55\u5934{fmt_number(shantou['energy'])}\u4e07\u3001"
        f"\u6d77\u95e8{fmt_number(haimen['energy'])}\u4e07\u3001"
        f"\u4e1c\u839e{fmt_number(dongguan['energy'])}\u4e07\u3001"
        f"\u6d77\u4e0a\u98ce\u7535{fmt_number(offshore['energy'])}\u4e07\u3001"
        f"\u9b80\u83b2{fmt_number(tuolian['energy'])}\u4e07\u3001"
        f"\u5f52\u6e56{fmt_number(guihu['energy'])}\u4e07\u5343\u74e6\u65f6\uff1b"
        f"\u5206\u516c\u53f8\u65e5\u524d\u5747\u4ef7{fmt_number(summary['averagePrice'])}\u5398/\u5343\u74e6\u65f6\uff0c"
        f"\u5176\u4e2d\u6c55\u5934{fmt_number(shantou['price'])}\u5398\u3001"
        f"\u6d77\u95e8{fmt_number(haimen['price'])}\u5398\u3001"
        f"\u4e1c\u839e{fmt_number(dongguan['price'])}\u5398\u3001"
        f"\u6d77\u4e0a\u98ce\u7535{fmt_number(offshore['price'])}\u5398\u3001"
        f"\u9b80\u83b2{fmt_number(tuolian['price'])}\u5398\u3001"
        f"\u5f52\u6e56{fmt_number(guihu['price'])}\u5398/\u5343\u74e6\u65f6\u3002"
    )


def replace_online_price_text(paragraph, price_summary: dict[str, Any]) -> None:
    text = paragraph.text
    companies = price_summary["companies"]
    shantou = companies["\u6c55\u5934"]
    haimen = companies["\u6d77\u95e8"]
    dongguan = companies["\u4e1c\u839e"]
    offshore = companies["\u6d77\u4e0a\u98ce\u7535"]
    tuolian = companies["\u9b80\u83b2"]
    guihu = companies["\u5f52\u6e56"]
    price_text = (
        f"\u5206\u516c\u53f8\u65e5\u524d\u5747\u4ef7{fmt_number(price_summary['averagePrice'])}\u5398/\u5343\u74e6\u65f6\uff0c"
        f"\u5176\u4e2d\u6c55\u5934{fmt_number(shantou)}\u5398\u3001"
        f"\u6d77\u95e8{fmt_number(haimen)}\u5398\u3001"
        f"\u4e1c\u839e{fmt_number(dongguan)}\u5398\u3001"
        f"\u6d77\u4e0a\u98ce\u7535{fmt_number(offshore)}\u5398\u3001"
        f"\u9b80\u83b2{fmt_number(tuolian)}\u5398\u3001"
        f"\u5f52\u6e56{fmt_number(guihu)}\u5398/\u5343\u74e6\u65f6\u3002"
    )
    updated = re.sub(
        r"\u5206\u516c\u53f8\u65e5\u524d\u5747\u4ef7.*?\u5f52\u6e56.*?\u5398/\u5343\u74e6\u65f6\u3002",
        price_text,
        text,
        count=1,
    )
    if updated != text:
        replace_paragraph(paragraph, updated)


def fill_forecast_table(table, forecast: dict[str, Any]) -> None:
    rows = forecast["rows"]
    for idx, item in enumerate(rows, start=1):
        replace_cell_text(table.cell(0, idx), date_label(item["forecastDate"]))
        replace_cell_text(table.cell(1, idx), str(int(round((item["dailyTotal"] or 0) / 10))))
    for hour in range(24):
        for col_idx, item in enumerate(rows, start=1):
            value = item["hours"][hour]
            replace_cell_text(table.cell(hour + 2, col_idx), "" if value is None else str(int(round(value / 10))))


def generate(
    report_date: str,
    template_path: Path,
    output_dir: Path,
    *,
    online_workbook: Path | None = None,
    day_ahead_workbook: Path | None = None,
    daily_clearing_workbook: Path | None = None,
    prefer_previous_template: bool = True,
) -> tuple[Path, Path]:
    plan = build_report_date_plan(report_date)
    actual_template_path = resolve_template_path(
        report_date,
        template_path,
        output_dir,
        prefer_previous_report=prefer_previous_template,
    )
    payload = run_daily_data(
        selected_date=plan["loadAndPriceSelectedDate"],
        compare_date=plan["loadAndPriceCompareDate"],
        subjects=LOAD_SUBJECT_CODES,
    )
    day_ahead_clearing = load_day_ahead_clearing_summary(plan["dayAheadDate"], day_ahead_workbook)
    online_prices = load_online_price_summary(plan["actualOnlineEnergyDate"], online_workbook)
    doc = Document(actual_template_path)
    replace_paragraph(doc.paragraphs[2], f"{day_number(plan['actualOnlineEnergyDate'])}\u65e5\uff1a")
    replace_paragraph(doc.paragraphs[6], f"{day_number(report_date)}\u65e5\uff1a")
    # P13 is left unchanged until the online energy table source is confirmed.
    replace_leading_date(doc.paragraphs[12], plan["actualOnlineEnergyDate"])
    replace_online_price_text(doc.paragraphs[12], online_prices)
    clear_paragraph_red(doc.paragraphs[12])
    replace_paragraph(doc.paragraphs[13], build_day_ahead_clearing_text(plan["dayAheadDate"], day_ahead_clearing))
    replace_paragraph(doc.paragraphs[15], build_strategy_text(payload))
    replace_paragraph(doc.paragraphs[20], f"\u88681  {date_label(plan['dailyClearingDate'])}\u5206\u516c\u53f8\u65e5\u6e05\u5206\u60c5\u51b5")
    replace_paragraph(doc.paragraphs[24], build_day_ahead_market_text(payload))
    replace_paragraph(doc.paragraphs[26], build_actual_market_text(payload))
    weather_error = None
    try:
        replace_weather_prefix(doc.paragraphs[28], report_date)
    except Exception as exc:
        weather_error = str(exc)
        clear_paragraph_red(doc.paragraphs[28])
        log(f"天气自动更新失败，已保留模板原文：{weather_error}")

    if len(doc.tables) >= 2:
        fill_forecast_table(doc.tables[1], payload["rollingForecast"])

    output_dir.mkdir(parents=True, exist_ok=True)
    out_path = output_dir / report_filename(report_date)
    try:
        doc.save(out_path)
    except PermissionError:
        dt = datetime.strptime(report_date, "%Y-%m-%d")
        out_path = output_dir / (
            f"\u6bcf\u65e5\u751f\u4ea7\u7ecf\u8425\u60c5\u51b5\u6c47\u62a5\uff08{dt.month}\u6708{dt.day}\u65e5\uff09_"
            f"{datetime.now().strftime('%H%M%S')}.docx"
        )
        doc.save(out_path)

    audit_path = output_dir / f"red_report_audit_{report_date}.json"
    audit = {
        "reportDate": report_date,
        "datePlan": plan,
        "template": str(actual_template_path),
        "templateFallback": str(template_path) if actual_template_path != template_path else None,
        "output": str(out_path),
        "onlineWorkbook": str(online_workbook) if online_workbook else None,
        "onlinePriceWorkbook": online_prices["workbook"],
        "dayAheadClearingWorkbook": day_ahead_clearing["workbook"],
        "dailyClearingWorkbook": str(daily_clearing_workbook) if daily_clearing_workbook else None,
        "weatherError": weather_error,
        "knownGaps": [
            "P13 上网电量来源尚未确认，电量暂保留模板原文；日前均价来自报告日出清情况 Excel 的快报（日前）工作表。",
            "P14 日前出清电量和日前价格来自出清情况 Excel 的快报（日前）工作表；火电取 F/G，新能源按日前出清表头取 G/H。",
            "P29 未来三天最高负荷来源尚未确认，暂不替换。",
        ],
    }
    audit_path.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    return out_path, audit_path


def main() -> int:
    parser = argparse.ArgumentParser(description="按红字模板生成试跑版报告")
    parser.add_argument("report_date", nargs="?", help="报告日期 YYYY-MM-DD")
    parser.add_argument("--template", default=str(TEMPLATE_PATH), help="红字模板路径")
    parser.add_argument("--output-dir", default=str(OUTPUT_DIR), help="输出目录")
    parser.add_argument("--online-workbook", default="", help="前一日报表 Excel，预留给上网电量段落")
    parser.add_argument("--day-ahead-workbook", default="", help="报告日报表 Excel，用于日前出清电量和日前价格")
    parser.add_argument("--daily-clearing-workbook", default="", help="前两日报表 Excel，预留给日清分表")
    args = parser.parse_args()
    try:
        if not args.report_date:
            raise ValueError("请提供报告日期，例如：python generate_red_marked_report.py 2026-05-13")
        report_date = validate_date(args.report_date)
        template_was_provided = any(
            item == "--template" or item.startswith("--template=")
            for item in sys.argv[1:]
        )
        out_path, audit_path = generate(
            report_date,
            Path(args.template),
            Path(args.output_dir),
            online_workbook=Path(args.online_workbook) if args.online_workbook else None,
            day_ahead_workbook=Path(args.day_ahead_workbook) if args.day_ahead_workbook else None,
            daily_clearing_workbook=Path(args.daily_clearing_workbook) if args.daily_clearing_workbook else None,
            prefer_previous_template=not template_was_provided,
        )
        log(f"已生成: {out_path}")
        log(f"审计说明: {audit_path}")
        return 0
    except KeyboardInterrupt:
        print("\n已取消。")
        return 130
    except Exception as exc:
        print(f"\n失败：{exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
